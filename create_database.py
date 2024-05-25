import os
import shutil
from langchain.document_loaders import DirectoryLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores.chroma import Chroma
from docx import Document as DocxDocument
from dotenv import load_dotenv

load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

CHROMA_PATH = "chroma"
DATA_PATH = "data"

def main():
    generate_data_store()

def generate_data_store():
    documents = load_documents()
    if documents:  # Ensure documents are loaded before proceeding
        chunks = split_text(documents)
        save_to_chroma(chunks)

def load_documents():
    # Ensure the directory exists
    if not os.path.exists(DATA_PATH):
        print(f"Directory not found: {DATA_PATH}")
        return []

    # List all .docx files in the directory
    docx_files = [file for file in os.listdir(DATA_PATH) if file.endswith('.docx')]
    documents = []

    # Load each .docx file
    for docx_file in docx_files:
        doc = DocxDocument(os.path.join(DATA_PATH, docx_file))
        text = "\n".join([para.text for para in doc.paragraphs])
        documents.append(Document(page_content=text, metadata={"title": docx_file}))

    return documents

def split_text(documents: list[Document]):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=300,
        chunk_overlap=100,
        length_function=len,
        add_start_index=True,
    )
    chunks = text_splitter.split_documents(documents)
    print(f"Split {len(documents)} documents into {len(chunks)} chunks.")

    if chunks:  # Ensure there are chunks before accessing
        document = chunks[0]
        print(document.page_content)
        print(document.metadata)

    return chunks

def save_to_chroma(chunks: list[Document]):
    # Clear out the database first.
    if os.path.exists(CHROMA_PATH):
        shutil.rmtree(CHROMA_PATH)

    # Create a new DB from the documents.
    db = Chroma.from_documents(
        chunks, OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY), persist_directory=CHROMA_PATH
    )
    db.persist()
    print(f"Saved {len(chunks)} chunks to {CHROMA_PATH}.")

if __name__ == "__main__":
    main()
